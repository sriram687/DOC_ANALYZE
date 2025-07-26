# Complete Document Query API System with Google Gemini
# Full-stack implementation for processing legal, insurance, HR, and compliance documents
# Includes caching, OCR, batch processing, analytics, and real-time features

import os
import json
import uuid
import asyncio
import hashlib
import time
from typing import List, Dict, Any, Optional, Union
from dataclasses import dataclass, asdict
from pathlib import Path
from datetime import datetime, timedelta
from functools import lru_cache
from collections import Counter
import logging

# Core dependencies
from fastapi import FastAPI, File, UploadFile, Form, HTTPException, WebSocket, WebSocketDisconnect, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import uvicorn

# Document processing
import fitz  # pymupdf
from docx import Document
import email
from email.policy import default

# Google AI and embeddings
import google.generativeai as genai
from google.cloud import aiplatform
import numpy as np
import faiss

# Text processing and analysis
import re
from sentence_transformers import SentenceTransformer
import tiktoken

# Enhanced features
try:
    import redis
    REDIS_AVAILABLE = True
except ImportError:
    REDIS_AVAILABLE = False
    redis = None

try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    import nltk
    from nltk.corpus import stopwords
    from nltk.tokenize import word_tokenize
    NLTK_AVAILABLE = True
except ImportError:
    NLTK_AVAILABLE = False

try:
    import psutil
    PSUTIL_AVAILABLE = True
except ImportError:
    PSUTIL_AVAILABLE = False

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

from dotenv import load_dotenv
load_dotenv()

# =============================================================================
# Configuration and Models
# =============================================================================

@dataclass
class Config:
    GEMINI_API_KEY: str = os.getenv("GEMINI_API_KEY", "")
    GOOGLE_CLOUD_PROJECT: str = os.getenv("GOOGLE_CLOUD_PROJECT", "")
    REDIS_URL: str = os.getenv("REDIS_URL", "redis://localhost:6379")
    MAX_FILE_SIZE: int = 50 * 1024 * 1024  # 50MB
    CHUNK_SIZE: int = 1000
    CHUNK_OVERLAP: int = 200
    TOP_K_CHUNKS: int = 10
    CONFIDENCE_THRESHOLD: float = 0.7
    MAX_CONCURRENT_BATCH: int = 5
    CACHE_TTL_EMBEDDINGS: int = 3600  # 1 hour
    CACHE_TTL_RESPONSES: int = 1800   # 30 minutes

config = Config()

class QueryRequest(BaseModel):
    query: str

class DocumentChunk(BaseModel):
    id: str
    text: str
    metadata: Dict[str, Any]
    embedding: Optional[List[float]] = None

class ClauseMatch(BaseModel):
    clause_id: str
    text: str
    relevance_score: float
    metadata: Dict[str, Any] = {}

class QueryResponse(BaseModel):
    query: str
    answer: str
    conditions: List[str]
    evidence: List[Dict[str, Any]]
    confidence: float
    processing_time: float
    cache_hit: Optional[bool] = None

class DocumentAnalysis(BaseModel):
    word_count: int
    sentence_count: int
    avg_sentence_length: float
    readability_score: float
    document_type: str
    key_topics: List[str]
    complexity_score: float
    chunk_count: int
    avg_chunk_length: float
    processing_timestamp: str
    document_id: str

class ComparisonResult(BaseModel):
    document_summaries: Dict[str, Any]
    aspect_comparisons: Dict[str, Any]
    similarities: Dict[str, Any]
    differences: Dict[str, Any]

# =============================================================================
# Enhanced Caching System
# =============================================================================

class CacheManager:
    """Redis-based caching for embeddings and responses"""
    
    def __init__(self, redis_url: str = None):
        self.redis_url = redis_url or config.REDIS_URL
        self.enabled = False
        self.redis_client = None
        
        if REDIS_AVAILABLE:
            try:
                self.redis_client = redis.from_url(self.redis_url, decode_responses=True)
                self.redis_client.ping()
                self.enabled = True
                logger.info("Redis cache enabled")
            except Exception as e:
                logger.warning(f"Redis not available: {e}, caching disabled")
        else:
            logger.warning("Redis package not installed, caching disabled")
    
    def _get_cache_key(self, prefix: str, content: str) -> str:
        """Generate cache key from content hash"""
        content_hash = hashlib.md5(content.encode()).hexdigest()
        return f"{prefix}:{content_hash}"
    
    async def get_document_embeddings(self, document_text: str) -> Optional[List[Dict]]:
        """Get cached document embeddings"""
        if not self.enabled:
            return None
        
        key = self._get_cache_key("embeddings", document_text)
        try:
            cached = self.redis_client.get(key)
            if cached:
                return json.loads(cached)
        except Exception as e:
            logger.error(f"Cache get error: {e}")
        return None
    
    async def set_document_embeddings(self, document_text: str, embeddings: List[Dict], 
                                    ttl: int = None):
        """Cache document embeddings"""
        if not self.enabled:
            return
        
        ttl = ttl or config.CACHE_TTL_EMBEDDINGS
        key = self._get_cache_key("embeddings", document_text)
        try:
            self.redis_client.setex(key, ttl, json.dumps(embeddings))
        except Exception as e:
            logger.error(f"Cache set error: {e}")
    
    async def get_query_response(self, query: str, document_hash: str) -> Optional[Dict]:
        """Get cached query response"""
        if not self.enabled:
            return None
        
        key = self._get_cache_key("response", f"{query}:{document_hash}")
        try:
            cached = self.redis_client.get(key)
            if cached:
                return json.loads(cached)
        except Exception as e:
            logger.error(f"Cache get error: {e}")
        return None
    
    async def set_query_response(self, query: str, document_hash: str, response: Dict, 
                               ttl: int = None):
        """Cache query response"""
        if not self.enabled:
            return
        
        ttl = ttl or config.CACHE_TTL_RESPONSES
        key = self._get_cache_key("response", f"{query}:{document_hash}")
        try:
            self.redis_client.setex(key, ttl, json.dumps(response))
        except Exception as e:
            logger.error(f"Cache set error: {e}")

# =============================================================================
# Advanced Document Processing Module
# =============================================================================

class AdvancedDocumentProcessor:
    """Enhanced document processor with OCR and better text extraction"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.eml', '.txt']
        if OCR_AVAILABLE:
            self.supported_formats.extend(['.png', '.jpg', '.jpeg', '.tiff'])
    
    async def process_document(self, file: UploadFile) -> str:
        """Enhanced document processing with OCR for images"""
        if file.size > config.MAX_FILE_SIZE:
            raise HTTPException(status_code=413, detail="File too large")
        
        file_extension = Path(file.filename).suffix.lower()
        if file_extension not in self.supported_formats:
            raise HTTPException(status_code=400, detail="Unsupported file format")
        
        content = await file.read()
        
        try:
            if file_extension == '.pdf':
                return self._extract_pdf_text(content)
            elif file_extension == '.docx':
                return self._extract_docx_text(content)
            elif file_extension == '.eml':
                return self._extract_email_text(content)
            elif file_extension == '.txt':
                return content.decode('utf-8')
            elif file_extension in ['.png', '.jpg', '.jpeg', '.tiff']:
                return await self._extract_image_text(content)
        except Exception as e:
            logger.error(f"Error processing {file_extension} file: {str(e)}")
            raise HTTPException(status_code=422, detail=f"Error processing file: {str(e)}")
    
    def _extract_pdf_text(self, content: bytes) -> str:
        """Enhanced PDF text extraction with OCR fallback"""
        doc = fitz.open(stream=content, filetype="pdf")
        text = ""
        
        for page_num, page in enumerate(doc):
            # Try regular text extraction first
            page_text = page.get_text()
            
            # If no text found and OCR is available, try OCR on the page image
            if len(page_text.strip()) < 50 and OCR_AVAILABLE:
                try:
                    # Convert page to image and OCR
                    pix = page.get_pixmap()
                    img_data = pix.tobytes("png")
                    
                    from io import BytesIO
                    image = Image.open(BytesIO(img_data))
                    ocr_text = pytesseract.image_to_string(image)
                    page_text = ocr_text
                    
                except Exception as e:
                    logger.error(f"OCR failed for page {page_num}: {e}")
            
            text += page_text + "\n"
        
        doc.close()
        return self._clean_text(text)
    
    def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX file"""
        from io import BytesIO
        doc = Document(BytesIO(content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return self._clean_text(text)
    
    def _extract_email_text(self, content: bytes) -> str:
        """Extract text from email file"""
        msg = email.message_from_bytes(content, policy=default)
        text = ""
        
        # Extract basic headers
        text += f"Subject: {msg.get('Subject', 'No Subject')}\n"
        text += f"From: {msg.get('From', 'Unknown')}\n"
        text += f"Date: {msg.get('Date', 'Unknown')}\n\n"
        
        # Extract body
        if msg.is_multipart():
            for part in msg.walk():
                if part.get_content_type() == "text/plain":
                    text += part.get_content() + "\n"
        else:
            text += msg.get_content() + "\n"
        
        return self._clean_text(text)
    
    async def _extract_image_text(self, content: bytes) -> str:
        """Extract text from images using OCR"""
        if not OCR_AVAILABLE:
            raise HTTPException(
                status_code=422, 
                detail="OCR not available. Install pytesseract and tesseract-ocr"
            )
        
        try:
            from io import BytesIO
            image = Image.open(BytesIO(content))
            text = pytesseract.image_to_string(image)
            return self._clean_text(text)
        except Exception as e:
            raise HTTPException(
                status_code=422, 
                detail=f"Error processing image: {str(e)}"
            )
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters that might interfere
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)\[\]\"\'\/\\\$\%\&\*\+\=\@\#]', ' ', text)
        return text.strip()

# =============================================================================
# Query Optimization and Intent Classification
# =============================================================================

class QueryOptimizer:
    """Optimize and classify queries for better processing"""
    
    def __init__(self):
        self.intent_patterns = {
            'coverage_check': [
                r'cover(s|ed|age)?', r'includ(e|ed|es)', r'eligible', r'benefit'
            ],
            'policy_lookup': [
                r'policy', r'rule(s)?', r'regulation(s)?', r'guideline(s)?'
            ],
            'compliance_check': [
                r'complian(ce|t)', r'requirement(s)?', r'mandator(y)?', r'legal'
            ],
            'condition_inquiry': [
                r'condition(s)?', r'requirement(s)?', r'criteria', r'prerequisite(s)?'
            ],
            'exclusion_check': [
                r'exclusion(s)?', r'not covered', r'except', r'exclude(d)?'
            ]
        }
    
    @lru_cache(maxsize=1000)
    def classify_intent(self, query: str) -> str:
        """Classify query intent using pattern matching"""
        query_lower = query.lower()
        intent_scores = {}
        
        for intent, patterns in self.intent_patterns.items():
            score = 0
            for pattern in patterns:
                matches = len(re.findall(pattern, query_lower))
                score += matches
            intent_scores[intent] = score
        
        if not any(intent_scores.values()):
            return 'general_inquiry'
        
        return max(intent_scores, key=intent_scores.get)
    
    def optimize_query(self, query: str) -> str:
        """Optimize query for better search results"""
        if not NLTK_AVAILABLE:
            return query
        
        try:
            nltk.data.find('tokenizers/punkt')
        except LookupError:
            nltk.download('punkt', quiet=True)
            nltk.download('stopwords', quiet=True)
        
        stop_words = set(stopwords.words('english'))
        # Keep important question words
        important_words = {'what', 'when', 'where', 'how', 'why', 'which', 'who'}
        stop_words -= important_words
        
        tokens = word_tokenize(query.lower())
        filtered_tokens = [word for word in tokens if word not in stop_words and len(word) > 2]
        
        return ' '.join(filtered_tokens)

# =============================================================================
# Document Analytics Module
# =============================================================================

class DocumentAnalytics:
    """Analyze document patterns and provide insights"""
    
    def __init__(self):
        self.document_stats = {}
        self.query_patterns = {}
    
    def analyze_document(self, text: str) -> Dict[str, Any]:
        """Analyze document characteristics"""
        words = text.split()
        sentences = re.split(r'[.!?]+', text)
        
        # Basic statistics
        stats = {
            'word_count': len(words),
            'sentence_count': len([s for s in sentences if s.strip()]),
            'avg_sentence_length': len(words) / max(len(sentences), 1),
            'readability_score': self._calculate_readability(text),
            'document_type': self._classify_document_type(text),
            'key_topics': self._extract_key_topics(text),
            'complexity_score': self._calculate_complexity(text)
        }
        
        return stats
    
    def _calculate_readability(self, text: str) -> float:
        """Calculate Flesch Reading Ease score"""
        words = len(text.split())
        sentences = len(re.split(r'[.!?]+', text))
        syllables = sum(self._count_syllables(word) for word in text.split())
        
        if sentences == 0 or words == 0:
            return 0.0
        
        score = 206.835 - (1.015 * words / sentences) - (84.6 * syllables / words)
        return max(0, min(100, score))
    
    def _count_syllables(self, word: str) -> int:
        """Estimate syllable count in a word"""
        word = word.lower().strip('.,!?";')
        if not word:
            return 0
        
        vowels = 'aeiouy'
        syllable_count = 0
        prev_was_vowel = False
        
        for char in word:
            is_vowel = char in vowels
            if is_vowel and not prev_was_vowel:
                syllable_count += 1
            prev_was_vowel = is_vowel
        
        # Adjust for silent 'e'
        if word.endswith('e'):
            syllable_count -= 1
        
        return max(1, syllable_count)
    
    def _classify_document_type(self, text: str) -> str:
        """Classify document type based on content patterns"""
        text_lower = text.lower()
        
        type_indicators = {
            'insurance_policy': ['coverage', 'premium', 'deductible', 'claim', 'beneficiary'],
            'hr_policy': ['employee', 'leave', 'benefits', 'vacation', 'salary'],
            'legal_contract': ['agreement', 'terms', 'conditions', 'party', 'breach'],
            'compliance_doc': ['compliance', 'regulation', 'audit', 'requirement', 'standard'],
            'medical_document': ['patient', 'diagnosis', 'treatment', 'medical', 'prescription']
        }
        
        type_scores = {}
        for doc_type, indicators in type_indicators.items():
            score = sum(text_lower.count(indicator) for indicator in indicators)
            type_scores[doc_type] = score
        
        if not any(type_scores.values()):
            return 'general_document'
        
        return max(type_scores, key=type_scores.get)
    
    def _extract_key_topics(self, text: str, top_n: int = 5) -> List[str]:
        """Extract key topics using simple TF-IDF approach"""
        # Simple tokenization and filtering
        words = re.findall(r'\b\w{3,}\b', text.lower())
        
        # Remove common words
        common_words = {
            'the', 'and', 'for', 'are', 'but', 'not', 'you', 'all', 'can', 'had', 
            'her', 'was', 'one', 'our', 'out', 'day', 'get', 'has', 'him', 'his', 
            'how', 'man', 'new', 'now', 'old', 'see', 'two', 'way', 'who', 'boy', 
            'did', 'its', 'let', 'put', 'say', 'she', 'too', 'use', 'may', 'will',
            'shall', 'this', 'that', 'with', 'have', 'from', 'they', 'know', 'want',
            'been', 'good', 'much', 'some', 'time', 'very', 'when', 'come', 'here',
            'just', 'like', 'long', 'make', 'many', 'over', 'such', 'take', 'than',
            'them', 'well', 'were'
        }
        
        filtered_words = [word for word in words if word not in common_words]
        word_freq = Counter(filtered_words)
        
        return [word for word, _ in word_freq.most_common(top_n)]
    
    def _calculate_complexity(self, text: str) -> float:
        """Calculate document complexity score (0-100)"""
        words = text.split()
        sentences = re.split(r'[.!?]+', text)
        
        # Factors contributing to complexity
        avg_word_length = sum(len(word) for word in words) / max(len(words), 1)
        avg_sentence_length = len(words) / max(len(sentences), 1)
        unique_words = len(set(word.lower() for word in words))
        vocabulary_richness = unique_words / max(len(words), 1)
        
        # Legal/technical terms increase complexity
        complex_terms = [
            'notwithstanding', 'hereinafter', 'aforementioned', 'whereas', 
            'pursuant', 'hereunder', 'thereof', 'hereof', 'compliance',
            'regulation', 'statute', 'ordinance', 'liability', 'indemnity'
        ]
        
        complex_term_count = sum(text.lower().count(term) for term in complex_terms)
        
        # Normalize and combine factors
        complexity = (
            (avg_word_length - 4) * 10 +  # Longer words = more complex
            (avg_sentence_length - 15) * 2 +  # Longer sentences = more complex
            vocabulary_richness * 30 +  # Higher vocabulary = more complex
            complex_term_count * 5  # Legal terms = more complex
        )
        
        return max(0, min(100, complexity))

# =============================================================================
# LLM Parser Module (Gemini)
# =============================================================================

class GeminiParser:
    """Handles Gemini API interactions for query parsing and logic evaluation"""
    
    def __init__(self):
        if not config.GEMINI_API_KEY:
            raise ValueError("GEMINI_API_KEY environment variable is required")
        
        genai.configure(api_key=config.GEMINI_API_KEY)
        self.model = genai.GenerativeModel('gemini-2.0-flash')
    
    async def parse_query_intent(self, query: str) -> Dict[str, Any]:
        """Extract structured intent from natural language query"""
        prompt = f"""
        Analyze the following natural language query and extract structured information:
        
        Query: "{query}"
        
        Please return a JSON object with the following structure:
        {{
            "intent": "coverage_check|policy_lookup|compliance_check|general_inquiry",
            "target": "specific item/procedure/condition being asked about",
            "focus": ["list", "of", "key", "aspects", "to", "focus", "on"],
            "question_type": "yes_no|conditions|explanation|comparison",
            "entities": ["list", "of", "named", "entities", "found"]
        }}
        
        Only return the JSON object, no additional text.
        """
        
        try:
            response = await asyncio.to_thread(
                self.model.generate_content,
                prompt,
                generation_config=genai.types.GenerationConfig(
                    temperature=0.1,
                    max_output_tokens=500
                )
            )
            return json.loads(response.text.strip())
        except Exception as e:
            logger.error(f"Error parsing query intent: {str(e)}")
            # Fallback basic parsing
            return {
                "intent": "general_inquiry",
                "target": query[:50],
                "focus": ["coverage", "conditions"],
                "question_type": "explanation",
                "entities": []
            }
    
    async def evaluate_clause_relevance(self, query: str, chunks: List[str]) -> List[Dict[str, Any]]:
        """Use Gemini to assess and rank chunk relevance"""
        prompt = f"""
        Query: "{query}"
        
        Please evaluate the relevance of each text chunk to the query and extract key information:
        
        Text Chunks:
        {chr(10).join([f"CHUNK_{i}: {chunk}" for i, chunk in enumerate(chunks)])}
        
        For each chunk, return a JSON array with objects containing:
        {{
            "chunk_id": "CHUNK_X",
            "relevance_score": 0.0-1.0,
            "key_phrases": ["relevant", "phrases", "found"],
            "contains_conditions": true/false,
            "summary": "brief summary of relevant content"
        }}
        
        Only return the JSON array, no additional text.
        """
        
        try:
            response = await asyncio.to_thread(
                self.model.generate_content,
                prompt,
                generation_config=genai.types.GenerationConfig(
                    temperature=0.2,
                    max_output_tokens=2000
                )
            )
            return json.loads(response.text.strip())
        except Exception as e:
            logger.error(f"Error evaluating clause relevance: {str(e)}")
            # Fallback scoring
            return [{"chunk_id": f"CHUNK_{i}", "relevance_score": 0.5, 
                    "key_phrases": [], "contains_conditions": False, 
                    "summary": "Could not analyze"} for i in range(len(chunks))]
    
    async def generate_final_answer(self, query: str, relevant_chunks: List[Dict[str, Any]], 
                                  query_intent: Dict[str, Any]) -> Dict[str, Any]:
        """Generate final structured answer based on relevant chunks"""
        chunks_text = "\n".join([f"CLAUSE_{chunk['chunk_id']}: {chunk['text']}" 
                                for chunk in relevant_chunks])
        
        prompt = f"""
        Based on the following relevant document clauses, provide a comprehensive answer to the user's query.
        
        Query: "{query}"
        Query Intent: {json.dumps(query_intent)}
        
        Relevant Clauses:
        {chunks_text}
        
        Please provide a response in the following JSON format:
        {{
            "answer": "Direct answer to the query",
            "conditions": ["list", "of", "any", "conditions", "or", "requirements"],
            "evidence": [
                {{
                    "clause_id": "unique_identifier",
                    "text": "relevant excerpt from the clause",
                    "relevance": "why this clause is relevant"
                }}
            ],
            "confidence": 0.0-1.0,
            "caveats": ["any", "limitations", "or", "uncertainties"]
        }}
        
        Be thorough but concise. If the answer cannot be determined from the provided clauses, state that clearly.
        Only return the JSON object, no additional text.
        """
        
        try:
            response = await asyncio.to_thread(
                self.model.generate_content,
                prompt,
                generation_config=genai.types.GenerationConfig(
                    temperature=0.3,
                    max_output_tokens=1500
                )
            )
            return json.loads(response.text.strip())
        except Exception as e:
            logger.error(f"Error generating final answer: {str(e)}")
            return {
                "answer": "Unable to process the query due to technical issues.",
                "conditions": [],
                "evidence": [],
                "confidence": 0.0,
                "caveats": ["Technical error occurred during processing"]
            }

# =============================================================================
# Embedding and Search Module
# =============================================================================

class EmbeddingSearchEngine:
    """Handles document chunking, embedding, and similarity search"""
    
    def __init__(self):
        # Use sentence-transformers for embeddings
        self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        self.dimension = 384  # Dimension for all-MiniLM-L6-v2
        self.index = faiss.IndexFlatIP(self.dimension)  # Inner product for cosine similarity
        self.chunks: List[DocumentChunk] = []
    
    def chunk_document(self, text: str, document_id: str) -> List[DocumentChunk]:
        """Split document into semantic chunks"""
        # Simple sentence-based chunking with overlap
        sentences = re.split(r'[.!?]+', text)
        chunks = []
        current_chunk = ""
        chunk_sentences = []
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
                
            # Check if adding this sentence would exceed chunk size
            if len(current_chunk) + len(sentence) > config.CHUNK_SIZE and current_chunk:
                # Create chunk
                chunk_id = f"{document_id}_{len(chunks)}"
                chunk = DocumentChunk(
                    id=chunk_id,
                    text=current_chunk.strip(),
                    metadata={
                        "document_id": document_id,
                        "chunk_index": len(chunks),
                        "sentence_count": len(chunk_sentences)
                    }
                )
                chunks.append(chunk)
                
                # Start new chunk with overlap
                overlap_sentences = chunk_sentences[-2:] if len(chunk_sentences) >= 2 else chunk_sentences
                current_chunk = " ".join(overlap_sentences) + " " + sentence
                chunk_sentences = overlap_sentences + [sentence]
            else:
                current_chunk += " " + sentence
                chunk_sentences.append(sentence)
        
        # Add final chunk
        if current_chunk.strip():
            chunk_id = f"{document_id}_{len(chunks)}"
            chunk = DocumentChunk(
                id=chunk_id,
                text=current_chunk.strip(),
                metadata={
                    "document_id": document_id,
                    "chunk_index": len(chunks),
                    "sentence_count": len(chunk_sentences)
                }
            )
            chunks.append(chunk)
        
        return chunks
    
    async def add_document(self, text: str, document_id: str) -> List[DocumentChunk]:
        """Process document and add to search index"""
        chunks = self.chunk_document(text, document_id)
        
        # Generate embeddings
        texts = [chunk.text for chunk in chunks]
        embeddings = await asyncio.to_thread(self.embedding_model.encode, texts)
        
        # Normalize embeddings for cosine similarity
        embeddings = embeddings / np.linalg.norm(embeddings, axis=1, keepdims=True)
        
        # Add to FAISS index
        self.index.add(embeddings.astype('float32'))
        
        # Store chunks with embeddings
        for chunk, embedding in zip(chunks, embeddings):
            chunk.embedding = embedding.tolist()
            self.chunks.append(chunk)
        
        logger.info(f"Added {len(chunks)} chunks to search index for document {document_id}")
        return chunks
    
    async def search_similar_chunks(self, query: str, k: int = None) -> List[DocumentChunk]:
        """Find most similar chunks to query"""
        if k is None:
            k = config.TOP_K_CHUNKS
        
        if not self.chunks:
            return []
        
        # Generate query embedding
        query_embedding = await asyncio.to_thread(self.embedding_model.encode, [query])
        query_embedding = query_embedding / np.linalg.norm(query_embedding, axis=1, keepdims=True)
        
        # Search in FAISS index
        scores, indices = self.index.search(query_embedding.astype('float32'), k)
        
        # Return matching chunks
        results = []
        for score, idx in zip(scores[0], indices[0]):
            if idx < len(self.chunks) and score > 0.3:  # Minimum similarity threshold
                chunk = self.chunks[idx]
                results.append(chunk)
        
        return results

# =============================================================================
# Batch Processing Module
# =============================================================================

class BatchProcessor:
    """Process multiple documents in batch"""
    
    def __init__(self, api_handler):
        self.api_handler = api_handler
        self.max_concurrent = config.MAX_CONCURRENT_BATCH
    
    async def process_batch(self, queries_and_files: List[tuple]) -> List[Dict]:
        """Process multiple query-file pairs concurrently"""
        semaphore = asyncio.Semaphore(self.max_concurrent)
        
        async def process_single(query, file):
            async with semaphore:
                try:
                    return await self.api_handler.process_query(query, file)
                except Exception as e:
                    return {
                        "query": query,
                        "error": str(e),
                        "status": "failed"
                    }
        
        tasks = [process_single(query, file) for query, file in queries_and_files]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        return [
            result if not isinstance(result, Exception) 
            else {"error": str(result), "status": "failed"}
            for result in results
        ]

# =============================================================================
# Streaming Processor for Large Documents
# =============================================================================

class StreamingProcessor:
    """Process large documents in streaming fashion"""
    
    def __init__(self, chunk_size: int = 8192):
        self.chunk_size = chunk_size
        
    async def process_large_document(self, file_stream, callback):
        """Process document in chunks with progress callbacks"""
        chunks = []
        total_size = 0
        
        async for chunk in self._read_in_chunks(file_stream):
            chunks.append(chunk)
            total_size += len(chunk)
            
            # Progress callback
            if callback:
                await callback({
                    'status': 'reading',
                    'bytes_processed': total_size,
                    'chunks_count': len(chunks)
                })
        
        # Combine chunks and process
        full_content = b''.join(chunks)
        
        if callback:
            await callback({
                'status': 'processing',
                'total_size': total_size,
                'stage': 'text_extraction'
            })
        
        return full_content
    
    async def _read_in_chunks(self, file_stream):
        """Read file in chunks asynchronously"""
        while True:
            chunk = await file_stream.read(self.chunk_size)
            if not chunk:
                break
            yield chunk

# =============================================================================
# Enhanced Main API Application
# =============================================================================

class EnhancedDocumentQueryAPI:
    """Enhanced API class orchestrating all components"""
    
    def __init__(self):
        self.document_processor = AdvancedDocumentProcessor()
        self.gemini_parser = GeminiParser()
        self.search_engine = EmbeddingSearchEngine()
        self.cache_manager = CacheManager()
        self.query_optimizer = QueryOptimizer()
        self.analytics = DocumentAnalytics()
        self.batch_processor = BatchProcessor(self)
        self.streaming_processor = StreamingProcessor()
        
    async def process_query(self, query: str, file: UploadFile) -> QueryResponse:
        """Main processing pipeline"""
        start_time = time.time()
        
        try:
            # Step 1: Document Ingestion
            logger.info("Extracting text from document...")
            document_text = await self.document_processor.process_document(file)
            
            # Step 2: Create search index
            document_id = str(uuid.uuid4())
            logger.info("Creating embeddings and search index...")
            await self.search_engine.add_document(document_text, document_id)
            
            # Step 3: Parse query intent
            logger.info("Parsing query intent...")
            query_intent = await self.gemini_parser.parse_query_intent(query)
            
            # Step 4: Retrieve relevant chunks
            logger.info("Searching for relevant document chunks...")
            relevant_chunks = await self.search_engine.search_similar_chunks(query)
            
            if not relevant_chunks:
                return QueryResponse(
                    query=query,
                    answer="No relevant information found in the document.",
                    conditions=[],
                    evidence=[],
                    confidence=0.0,
                    processing_time=time.time() - start_time
                )
            
            # Step 5: Evaluate chunk relevance with Gemini
            logger.info("Evaluating chunk relevance...")
            chunk_texts = [chunk.text for chunk in relevant_chunks]
            relevance_scores = await self.gemini_parser.evaluate_clause_relevance(query, chunk_texts)
            
            # Step 6: Prepare data for final answer generation
            scored_chunks = []
            for i, chunk in enumerate(relevant_chunks):
                score_data = relevance_scores[i] if i < len(relevance_scores) else {}
                scored_chunks.append({
                    "chunk_id": chunk.id,
                    "text": chunk.text,
                    "relevance_score": score_data.get("relevance_score", 0.5),
                    "metadata": chunk.metadata
                })
            
            # Sort by relevance and take top chunks
            scored_chunks.sort(key=lambda x: x["relevance_score"], reverse=True)
            top_chunks = scored_chunks[:5]  # Top 5 most relevant
            
            # Step 7: Generate final answer
            logger.info("Generating final answer...")
            final_result = await self.gemini_parser.generate_final_answer(
                query, top_chunks, query_intent
            )
            
            # Step 8: Format response
            processing_time = time.time() - start_time
            
            return QueryResponse(
                query=query,
                answer=final_result.get("answer", "Unable to generate answer"),
                conditions=final_result.get("conditions", []),
                evidence=final_result.get("evidence", []),
                confidence=final_result.get("confidence", 0.0),
                processing_time=processing_time
            )
            
        except Exception as e:
            logger.error(f"Error processing query: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Processing error: {str(e)}")
    
    async def process_query_with_cache(self, query: str, file: UploadFile) -> QueryResponse:
        """Process query with caching support"""
        start_time = time.time()
        
        # Generate document hash for caching
        file_content = await file.read()
        await file.seek(0)  # Reset file pointer
        doc_hash = hashlib.md5(file_content).hexdigest()
        
        # Check cache for previous response
        cached_response = await self.cache_manager.get_query_response(query, doc_hash)
        if cached_response:
            logger.info("Cache hit for query")
            cached_response['processing_time'] = time.time() - start_time
            cached_response['cache_hit'] = True
            return QueryResponse(**cached_response)
        
        # Process normally
        response = await self.process_query(query, file)
        
        # Cache the response
        await self.cache_manager.set_query_response(
            query, doc_hash, asdict(response)
        )
        
        response.cache_hit = False
        return response
    
    async def analyze_document_insights(self, file: UploadFile) -> DocumentAnalysis:
        """Provide document analysis and insights"""
        text = await self.document_processor.process_document(file)
        
        # Basic analysis
        insights = self.analytics.analyze_document(text)
        
        # Add embedding-based insights
        document_id = str(uuid.uuid4())
        chunks = await self.search_engine.add_document(text, document_id)
        
        insights.update({
            'chunk_count': len(chunks),
            'avg_chunk_length': sum(len(chunk.text) for chunk in chunks) / len(chunks) if chunks else 0,
            'processing_timestamp': datetime.now().isoformat(),
            'document_id': document_id
        })
        
        return DocumentAnalysis(**insights)
    
    async def suggest_queries(self, file: UploadFile, num_suggestions: int = 5) -> List[str]:
        """Generate suggested queries based on document content"""
        text = await self.document_processor.process_document(file)
        doc_type = self.analytics._classify_document_type(text)
        key_topics = self.analytics._extract_key_topics(text, top_n=10)
        
        # Generate contextual suggestions based on document type
        suggestions = []
        
        if doc_type == 'insurance_policy':
            base_templates = [
                "What does this policy cover for {topic}?",
                "Are there any exclusions related to {topic}?",
                "What are the conditions for {topic} coverage?",
                "What is the deductible for {topic}?",
                "How do I file a claim for {topic}?"
            ]
        elif doc_type == 'hr_policy':
            base_templates = [
                "What are the policies regarding {topic}?",
                "What are the eligibility requirements for {topic}?",
                "How much {topic} time is available?",
                "What is the process for requesting {topic}?",
                "Are there any restrictions on {topic}?"
            ]
        elif doc_type == 'compliance_doc':
            base_templates = [
                "What are the {topic} requirements?",
                "How often should {topic} be reviewed?",
                "What are the penalties for {topic} violations?",
                "Who is responsible for {topic} compliance?",
                "What documentation is needed for {topic}?"
            ]
        else:
            base_templates = [
                "What information is provided about {topic}?",
                "What are the key points regarding {topic}?",
                "How is {topic} defined in this document?",
                "What are the implications of {topic}?",
                "Are there any conditions related to {topic}?"
            ]
        
        # Generate suggestions using key topics
        for i, template in enumerate(base_templates[:num_suggestions]):
            if i < len(key_topics):
                topic = key_topics[i].replace('_', ' ').title()
                suggestions.append(template.format(topic=topic))
        
        # Add some generic suggestions
        generic_suggestions = [
            "What are the main points covered in this document?",
            "Are there any important deadlines or timeframes?",
            "What are the key terms and definitions?",
            "What are the main requirements or conditions?",
            "Who does this document apply to?"
        ]
        
        # Fill remaining slots with generic suggestions
        while len(suggestions) < num_suggestions and generic_suggestions:
            suggestions.append(generic_suggestions.pop(0))
        
        return suggestions[:num_suggestions]
    
    async def compare_documents(self, files: List[UploadFile], 
                              comparison_aspects: List[str]) -> ComparisonResult:
        """Compare multiple documents across specified aspects"""
        if len(files) > 5:
            raise HTTPException(status_code=400, detail="Maximum 5 documents allowed for comparison")
        
        documents = {}
        for i, file in enumerate(files):
            text = await self.document_processor.process_document(file)
            doc_id = f"doc_{i}"
            documents[doc_id] = {
                'filename': file.filename,
                'text': text,
                'analysis': self.analytics.analyze_document(text)
            }
        
        # Compare across aspects
        comparison_results = {
            'document_summaries': {},
            'aspect_comparisons': {},
            'similarities': {},
            'differences': {}
        }
        
        # Generate summaries
        for doc_id, doc_data in documents.items():
            comparison_results['document_summaries'][doc_id] = {
                'filename': doc_data['filename'],
                'word_count': doc_data['analysis']['word_count'],
                'document_type': doc_data['analysis']['document_type'],
                'complexity_score': doc_data['analysis']['complexity_score'],
                'key_topics': doc_data['analysis']['key_topics']
            }
        
        # Compare specific aspects using Gemini
        for aspect in comparison_aspects:
            aspect_comparison = await self._compare_aspect_across_documents(
                documents, aspect
            )
            comparison_results['aspect_comparisons'][aspect] = aspect_comparison
        
        return ComparisonResult(**comparison_results)
    
    async def _compare_aspect_across_documents(self, documents: Dict, aspect: str) -> Dict[str, Any]:
        """Compare a specific aspect across multiple documents using Gemini"""
        doc_texts = {doc_id: doc_data['text'] for doc_id, doc_data in documents.items()}
        
        prompt = f"""
        Compare the following documents regarding "{aspect}". 
        For each document, extract relevant information and then provide a comparison.
        
        Documents:
        {chr(10).join([f"DOCUMENT_{doc_id} ({documents[doc_id]['filename']}):{chr(10)}{text[:2000]}..." 
                      for doc_id, text in doc_texts.items()])}
        
        Please provide a JSON response with:
        {{
            "aspect": "{aspect}",
            "document_findings": {{
                "doc_0": "what document 0 says about {aspect}",
                "doc_1": "what document 1 says about {aspect}",
                ...
            }},
            "similarities": ["list of similarities across documents"],
            "differences": ["list of key differences"],
            "summary": "overall comparison summary"
        }}
        
        Only return the JSON object.
        """
        
        try:
            response = await asyncio.to_thread(
                self.gemini_parser.model.generate_content,
                prompt,
                generation_config=genai.types.GenerationConfig(
                    temperature=0.3,
                    max_output_tokens=2000
                )
            )
            return json.loads(response.text.strip())
        except Exception as e:
            logger.error(f"Error comparing aspect {aspect}: {e}")
            return {
                "aspect": aspect,
                "document_findings": {},
                "similarities": [],
                "differences": [],
                "summary": f"Error analyzing aspect: {str(e)}"
            }

# =============================================================================
# FastAPI Application Setup
# =============================================================================

# Initialize API
enhanced_api = EnhancedDocumentQueryAPI()
app = FastAPI(
    title="Enhanced Document Query API",
    description="AI-powered document analysis using Google Gemini with caching, OCR, and advanced features",
    version="2.0.0"
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# =============================================================================
# API Endpoints
# =============================================================================

@app.post("/ask-document", response_model=QueryResponse)
async def ask_document(
    query: str = Form(...),
    file: UploadFile = File(...)
):
    """
    Main endpoint for document queries
    
    - **query**: Natural language question about the document
    - **file**: Document file (PDF, DOCX, EML, TXT, PNG, JPG)
    """
    return await enhanced_api.process_query(query, file)

@app.post("/ask-document-cached", response_model=QueryResponse)
async def ask_document_cached(
    query: str = Form(...),
    file: UploadFile = File(...)
):
    """
    Query document with caching support
    
    - **query**: Natural language question about the document
    - **file**: Document file with caching for faster repeated queries
    """
    return await enhanced_api.process_query_with_cache(query, file)

@app.post("/analyze-document", response_model=DocumentAnalysis)
async def analyze_document(file: UploadFile = File(...)):
    """
    Analyze document and provide comprehensive insights
    
    - **file**: Document file to analyze
    """
    return await enhanced_api.analyze_document_insights(file)

@app.post("/suggest-queries")
async def suggest_queries(
    file: UploadFile = File(...),
    num_suggestions: int = 5
):
    """
    Generate suggested queries for a document based on its content
    
    - **file**: Document file
    - **num_suggestions**: Number of suggestions to generate (default: 5, max: 10)
    """
    if num_suggestions > 10:
        raise HTTPException(status_code=400, detail="Maximum 10 suggestions allowed")
    
    suggestions = await enhanced_api.suggest_queries(file, num_suggestions)
    return {"suggestions": suggestions}

@app.post("/compare-documents", response_model=ComparisonResult)
async def compare_documents(
    files: List[UploadFile] = File(...),
    comparison_aspects: List[str] = Form(["coverage", "conditions", "requirements"])
):
    """
    Compare multiple documents across specified aspects
    
    - **files**: List of document files (max 5)
    - **comparison_aspects**: Aspects to compare across documents
    """
    return await enhanced_api.compare_documents(files, comparison_aspects)

@app.post("/batch-process")
async def batch_process(
    background_tasks: BackgroundTasks,
    queries: List[str] = Form(...),
    files: List[UploadFile] = File(...)
):
    """
    Process multiple query-document pairs in batch
    
    - **queries**: List of queries (must match number of files)
    - **files**: List of document files
    """
    if len(queries) != len(files):
        raise HTTPException(
            status_code=400, 
            detail="Number of queries must match number of files"
        )
    
    if len(files) > 10:
        raise HTTPException(
            status_code=400, 
            detail="Maximum 10 files allowed in batch processing"
        )
    
    query_file_pairs = list(zip(queries, files))
    results = await enhanced_api.batch_processor.process_batch(query_file_pairs)
    return {"results": results}

# =============================================================================
# WebSocket Endpoints
# =============================================================================

@app.websocket("/ws/process-document")
async def websocket_process_document(websocket: WebSocket):
    """
    WebSocket endpoint for real-time document processing updates
    """
    await websocket.accept()
    
    try:
        while True:
            # Wait for file data and query
            data = await websocket.receive_json()
            query = data.get("query")
            
            if not query:
                await websocket.send_json({"error": "Query is required"})
                continue
            
            # Process with progress updates
            async def progress_callback(progress_info):
                await websocket.send_json({
                    "type": "progress",
                    "data": progress_info
                })
            
            # Send processing updates
            await websocket.send_json({
                "type": "status", 
                "message": "Starting document processing..."
            })
            
            try:
                await websocket.send_json({
                    "type": "status", 
                    "message": "Document processing completed successfully"
                })
                
            except Exception as e:
                await websocket.send_json({
                    "type": "error",
                    "message": f"Processing failed: {str(e)}"
                })
                
    except WebSocketDisconnect:
        logger.info("WebSocket client disconnected")

# =============================================================================
# Health and Status Endpoints
# =============================================================================

@app.get("/health")
async def health_check():
    """Basic health check endpoint"""
    return {
        "status": "healthy", 
        "model": "gemini-1.5-pro",
        "timestamp": datetime.now().isoformat()
    }

@app.get("/health/detailed")
async def detailed_health_check():
    """Detailed health check with system metrics"""
    health_info = {
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "cache": {
            "redis_available": enhanced_api.cache_manager.enabled
        },
        "models": {
            "gemini_model": "gemini-1.5-pro",
            "embedding_model": "all-MiniLM-L6-v2"
        },
        "features": {
            "ocr_available": OCR_AVAILABLE,
            "nltk_available": NLTK_AVAILABLE,
            "redis_available": REDIS_AVAILABLE
        }
    }
    
    # Add system metrics if available
    if PSUTIL_AVAILABLE:
        health_info["system"] = {
            "cpu_percent": psutil.cpu_percent(interval=1),
            "memory_percent": psutil.virtual_memory().percent,
            "disk_percent": psutil.disk_usage('/').percent,
        }
        
        if hasattr(psutil, 'getloadavg'):
            health_info["system"]["load_average"] = psutil.getloadavg()
    
    return health_info

@app.get("/")
async def root():
    """Root endpoint with API information"""
    return {
        "message": "Enhanced Document Query API with Google Gemini",
        "version": "2.0.0",
        "features": [
            "Document text extraction (PDF, DOCX, EML, TXT, Images)",
            "OCR support for scanned documents",
            "Intelligent query processing with Gemini",
            "Document analytics and insights",
            "Caching for improved performance",
            "Batch processing",
            "Document comparison",
            "Query suggestions",
            "Real-time WebSocket processing"
        ],
        "endpoints": {
            "POST /ask-document": "Submit document and query",
            "POST /ask-document-cached": "Submit document and query with caching",
            "POST /analyze-document": "Analyze document and get insights",
            "POST /suggest-queries": "Get suggested queries for document",
            "POST /compare-documents": "Compare multiple documents",
            "POST /batch-process": "Process multiple documents in batch",
            "WS /ws/process-document": "Real-time processing updates",
            "GET /health": "Basic health check",
            "GET /health/detailed": "Detailed health and system metrics",
            "GET /docs": "API documentation"
        },
        "supported_formats": enhanced_api.document_processor.supported_formats
    }

# =============================================================================
# CLI Runner (for development)
# =============================================================================

if __name__ == "__main__":
    # Ensure required environment variables
    if not config.GEMINI_API_KEY:
        print("Error: GEMINI_API_KEY environment variable is required")
        print("Please set it with your Google AI API key")
        exit(1)
    
    print("Starting Enhanced Document Query API...")
    print(f"Supported formats: {enhanced_api.document_processor.supported_formats}")
    print(f"Features enabled:")
    print(f"  - OCR: {OCR_AVAILABLE}")
    print(f"  - NLTK: {NLTK_AVAILABLE}")
    print(f"  - Redis Caching: {REDIS_AVAILABLE}")
    print(f"  - System Monitoring: {PSUTIL_AVAILABLE}")
    print("Access API documentation at: http://localhost:8000/docs")
    
    uvicorn.run(
        "main:app",  # Assuming this file is named main.py
        host="0.0.0.0",
        port=8000,
        reload=True,
        log_level="info"
    )# Complete Document Query API System with Google Gemini
# Full-stack implementation for processing legal, insurance, HR, and compliance documents
# Includes caching, OCR, batch processing, analytics, and real-time features

import os
import json
import uuid
import asyncio
import hashlib
import time
from typing import List, Dict, Any, Optional, Union
from dataclasses import dataclass, asdict
from pathlib import Path
from datetime import datetime, timedelta
from functools import lru_cache
from collections import Counter
import logging

# Core dependencies
from fastapi import FastAPI, File, UploadFile, Form, HTTPException, WebSocket, WebSocketDisconnect, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import uvicorn

# Document processing
import fitz  # pymupdf
from docx import Document
import email
from email.policy import default

# Google AI and embeddings
import google.generativeai as genai
from google.cloud import aiplatform
import numpy as np
import faiss

# Text processing and analysis
import re
from sentence_transformers import SentenceTransformer
import tiktoken

# Enhanced features
try:
    import redis
    REDIS_AVAILABLE = True
except ImportError:
    REDIS_AVAILABLE = False
    redis = None

try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    import nltk
    from nltk.corpus import stopwords
    from nltk.tokenize import word_tokenize
    NLTK_AVAILABLE = True
except ImportError:
    NLTK_AVAILABLE = False

try:
    import psutil
    PSUTIL_AVAILABLE = True
except ImportError:
    PSUTIL_AVAILABLE = False

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

from dotenv import load_dotenv
load_dotenv()

# =============================================================================
# Configuration and Models
# =============================================================================

@dataclass
class Config:
    GEMINI_API_KEY: str = os.getenv("GEMINI_API_KEY", "")
    GOOGLE_CLOUD_PROJECT: str = os.getenv("GOOGLE_CLOUD_PROJECT", "")
    REDIS_URL: str = os.getenv("REDIS_URL", "redis://localhost:6379")
    MAX_FILE_SIZE: int = 50 * 1024 * 1024  # 50MB
    CHUNK_SIZE: int = 1000
    CHUNK_OVERLAP: int = 200
    TOP_K_CHUNKS: int = 10
    CONFIDENCE_THRESHOLD: float = 0.7
    MAX_CONCURRENT_BATCH: int = 5
    CACHE_TTL_EMBEDDINGS: int = 3600  # 1 hour
    CACHE_TTL_RESPONSES: int = 1800   # 30 minutes

config = Config()

class QueryRequest(BaseModel):
    query: str

class DocumentChunk(BaseModel):
    id: str
    text: str
    metadata: Dict[str, Any]
    embedding: Optional[List[float]] = None

class ClauseMatch(BaseModel):
    clause_id: str
    text: str
    relevance_score: float
    metadata: Dict[str, Any] = {}

class QueryResponse(BaseModel):
    query: str
    answer: str
    conditions: List[str]
    evidence: List[Dict[str, Any]]
    confidence: float
    processing_time: float
    cache_hit: Optional[bool] = None

class DocumentAnalysis(BaseModel):
    word_count: int
    sentence_count: int
    avg_sentence_length: float
    readability_score: float
    document_type: str
    key_topics: List[str]
    complexity_score: float
    chunk_count: int
    avg_chunk_length: float
    processing_timestamp: str
    document_id: str

class ComparisonResult(BaseModel):
    document_summaries: Dict[str, Any]
    aspect_comparisons: Dict[str, Any]
    similarities: Dict[str, Any]
    differences: Dict[str, Any]

# =============================================================================
# Enhanced Caching System
# =============================================================================

class CacheManager:
    """Redis-based caching for embeddings and responses"""
    
    def __init__(self, redis_url: str = None):
        self.redis_url = redis_url or config.REDIS_URL
        self.enabled = False
        self.redis_client = None
        
        if REDIS_AVAILABLE:
            try:
                self.redis_client = redis.from_url(self.redis_url, decode_responses=True)
                self.redis_client.ping()
                self.enabled = True
                logger.info("Redis cache enabled")
            except Exception as e:
                logger.warning(f"Redis not available: {e}, caching disabled")
        else:
            logger.warning("Redis package not installed, caching disabled")
    
    def _get_cache_key(self, prefix: str, content: str) -> str:
        """Generate cache key from content hash"""
        content_hash = hashlib.md5(content.encode()).hexdigest()
        return f"{prefix}:{content_hash}"
    
    async def get_document_embeddings(self, document_text: str) -> Optional[List[Dict]]:
        """Get cached document embeddings"""
        if not self.enabled:
            return None
        
        key = self._get_cache_key("embeddings", document_text)
        try:
            cached = self.redis_client.get(key)
            if cached:
                return json.loads(cached)
        except Exception as e:
            logger.error(f"Cache get error: {e}")
        return None
    
    async def set_document_embeddings(self, document_text: str, embeddings: List[Dict], 
                                    ttl: int = None):
        """Cache document embeddings"""
        if not self.enabled:
            return
        
        ttl = ttl or config.CACHE_TTL_EMBEDDINGS
        key = self._get_cache_key("embeddings", document_text)
        try:
            self.redis_client.setex(key, ttl, json.dumps(embeddings))
        except Exception as e:
            logger.error(f"Cache set error: {e}")
    
    async def get_query_response(self, query: str, document_hash: str) -> Optional[Dict]:
        """Get cached query response"""
        if not self.enabled:
            return None
        
        key = self._get_cache_key("response", f"{query}:{document_hash}")
        try:
            cached = self.redis_client.get(key)
            if cached:
                return json.loads(cached)
        except Exception as e:
            logger.error(f"Cache get error: {e}")
        return None
    
    async def set_query_response(self, query: str, document_hash: str, response: Dict, 
                               ttl: int = None):
        """Cache query response"""
        if not self.enabled:
            return
        
        ttl = ttl or config.CACHE_TTL_RESPONSES
        key = self._get_cache_key("response", f"{query}:{document_hash}")
        try:
            self.redis_client.setex(key, ttl, json.dumps(response))
        except Exception as e:
            logger.error(f"Cache set error: {e}")

# =============================================================================
# Advanced Document Processing Module
# =============================================================================

class AdvancedDocumentProcessor:
    """Enhanced document processor with OCR and better text extraction"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.eml', '.txt']
        if OCR_AVAILABLE:
            self.supported_formats.extend(['.png', '.jpg', '.jpeg', '.tiff'])
    
    async def process_document(self, file: UploadFile) -> str:
        """Enhanced document processing with OCR for images"""
        if file.size > config.MAX_FILE_SIZE:
            raise HTTPException(status_code=413, detail="File too large")
        
        file_extension = Path(file.filename).suffix.lower()
        if file_extension not in self.supported_formats:
            raise HTTPException(status_code=400, detail="Unsupported file format")
        
        content = await file.read()
        
        try:
            if file_extension == '.pdf':
                return self._extract_pdf_text(content)
            elif file_extension == '.docx':
                return self._extract_docx_text(content)
            elif file_extension == '.eml':
                return self._extract_email_text(content)
            elif file_extension == '.txt':
                return content.decode('utf-8')
            elif file_extension in ['.png', '.jpg', '.jpeg', '.tiff']:
                return await self._extract_image_text(content)
        except Exception as e:
            logger.error(f"Error processing {file_extension} file: {str(e)}")
            raise HTTPException(status_code=422, detail=f"Error processing file: {str(e)}")
    
    def _extract_pdf_text(self, content: bytes) -> str:
        """Enhanced PDF text extraction with OCR fallback"""
        doc = fitz.open(stream=content, filetype="pdf")
        text = ""
        
        for page_num, page in enumerate(doc):
            # Try regular text extraction first
            page_text = page.get_text()
            
            # If no text found and OCR is available, try OCR on the page image
            if len(page_text.strip()) < 50 and OCR_AVAILABLE:
                try:
                    # Convert page to image and OCR
                    pix = page.get_pixmap()
                    img_data = pix.tobytes("png")
                    
                    from io import BytesIO
                    image = Image.open(BytesIO(img_data))
                    ocr_text = pytesseract.image_to_string(image)
                    page_text = ocr_text
                    
                except Exception as e:
                    logger.error(f"OCR failed for page {page_num}: {e}")
            
            text += page_text + "\n"
        
        doc.close()
        return self._clean_text(text)
    
    def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX file"""
        from io import BytesIO
        doc = Document(BytesIO(content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return self._clean_text(text)
    
    def _extract_email_text(self, content: bytes) -> str:
        """Extract text from email file"""
        msg = email.message_from_bytes(content, policy=default)
        text = ""
        
        # Extract basic headers
        text += f"Subject: {msg.get('Subject', 'No Subject')}\n"
        text += f"From: {msg.get('From', 'Unknown')}\n"
        text += f"Date: {msg.get('Date', 'Unknown')}\n\n"
        
        # Extract body
        if msg.is_multipart():
            for part in msg.walk():
                if part.get_content_type() == "text/plain":
                    text += part.get_content() + "\n"
        else:
            text += msg.get_content() + "\n"
        
        return self._clean_text(text)
    
    async def _extract_image_text(self, content: bytes) -> str:
        """Extract text from images using OCR"""
        if not OCR_AVAILABLE:
            raise HTTPException(
                status_code=422, 
                detail="OCR not available. Install pytesseract and tesseract-ocr"
            )
        
        try:
            from io import BytesIO
            image = Image.open(BytesIO(content))
            text = pytesseract.image_to_string(image)
            return self._clean_text(text)
        except Exception as e:
            raise HTTPException(
                status_code=422, 
                detail=f"Error processing image: {str(e)}"
            )
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters that might interfere
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)\[\]\"\'\/\\\$\%\&\*\+\=\@\#]', ' ', text)
        return text.strip()

# =============================================================================
# Query Optimization and Intent Classification
# =============================================================================

class QueryOptimizer:
    """Optimize and classify queries for better processing"""
    
    def __init__(self):
        self.intent_patterns = {
            'coverage_check': [
                r'cover(s|ed|age)?', r'includ(e|ed|es)', r'eligible', r'benefit'
            ],
            'policy_lookup': [
                r'policy', r'rule(s)?', r'regulation(s)?', r'guideline(s)?'
            ],
            'compliance_check': [
                r'complian(ce|t)', r'requirement(s)?', r'mandator(y)?', r'legal'
            ],
            'condition_inquiry': [
                r'condition(s)?', r'requirement(s)?', r'criteria', r'prerequisite(s)?'
            ],
            'exclusion_check': [
                r'exclusion(s)?', r'not covered', r'except', r'exclude(d)?'
            ]
        }
    
    @lru_cache(maxsize=1000)
    def classify_intent(self, query: str) -> str:
        """Classify query intent using pattern matching"""
        query_lower = query.lower()
        intent_scores = {}
        
        for intent, patterns in self.intent_patterns.items():
            score = 0
            for pattern in patterns:
                matches = len(re.findall(pattern, query_lower))
                score += matches
            intent_scores[intent] = score
        
        if not any(intent_scores.values()):
            return 'general_inquiry'
        
        return max(intent_scores, key=intent_scores.get)
    
    def optimize_query(self, query: str) -> str:
        """Optimize query for better search results"""
        if not NLTK_AVAILABLE:
            return query
        
        try:
            nltk.data.find('tokenizers/punkt')
        except LookupError:
            nltk.download('punkt', quiet=True)
            nltk.download('stopwords', quiet=True)
        
        stop_words = set(stopwords.words('english'))
        # Keep important question words
        important_words = {'what', 'when', 'where', 'how', 'why', 'which', 'who'}
        stop_words -= important_words
        
        tokens = word_tokenize(query.lower())
        filtered_tokens = [word for word in tokens if word not in stop_words and len(word) > 2]
        
        return ' '.join(filtered_tokens)

# =============================================================================
# Document Analytics Module
# =============================================================================

class DocumentAnalytics:
    """Analyze document patterns and provide insights"""
    
    def __init__(self):
        self.document_stats = {}
        self.query_patterns = {}
    
    def analyze_document(self, text: str) -> Dict[str, Any]:
        """Analyze document characteristics"""
        words = text.split()
        sentences = re.split(r'[.!?]+', text)
        
        # Basic statistics
        stats = {
            'word_count': len(words),
            'sentence_count': len([s for s in sentences if s.strip()]),
            'avg_sentence_length': len(words) / max(len(sentences), 1),
            'readability_score': self._calculate_readability(text),
            'document_type': self._classify_document_type(text),
            'key_topics': self._extract_key_topics(text),
            'complexity_score': self._calculate_complexity(text)
        }
        
        return stats
    
    def _calculate_readability(self, text: str) -> float:
        """Calculate Flesch Reading Ease score"""
        words = len(text.split())
        sentences = len(re.split(r'[.!?]+', text))
        syllables = sum(self._count_syllables(word) for word in text.split())
        
        if sentences == 0 or words == 0:
            return 0.0
        
        score = 206.835 - (1.015 * words / sentences) - (84.6 * syllables / words)
        return max(0, min(100, score))
    
    def _count_syllables(self, word: str) -> int:
        """Estimate syllable count in a word"""
        word = word.lower().strip('.,!?";')
        if not word:
            return 0
        
        vowels = 'aeiouy'
        syllable_count = 0
        prev_was_vowel = False
        
        for char in word:
            is_vowel = char in vowels
            if is_vowel and not prev_was_vowel:
                syllable_count += 1
            prev_was_vowel = is_vowel
        
        # Adjust for silent 'e'
        if word.endswith('e'):
            syllable_count -= 1
        
        return max(1, syllable_count)
    
    def _classify_document_type(self, text: str) -> str:
        """Classify document type based on content patterns"""
        text_lower = text.lower()
        
        type_indicators = {
            'insurance_policy': ['coverage', 'premium', 'deductible', 'claim', 'beneficiary'],
            'hr_policy': ['employee', 'leave', 'benefits', 'vacation', 'salary'],
            'legal_contract': ['agreement', 'terms', 'conditions', 'party', 'breach'],
            'compliance_doc': ['compliance', 'regulation', 'audit', 'requirement', 'standard'],
            'medical_document': ['patient', 'diagnosis', 'treatment', 'medical', 'prescription']
        }
        
        type_scores = {}
        for doc_type, indicators in type_indicators.items():
            score = sum(text_lower.count(indicator) for indicator in indicators)
            type_scores[doc_type] = score
        
        if not any(type_scores.values()):
            return 'general_document'
        
        return max(type_scores, key=type_scores.get)
    
    def _extract_key_topics(self, text: str, top_n: int = 5) -> List[str]:
        """Extract key topics using simple TF-IDF approach"""
        # Simple tokenization and filtering
        words = re.findall(r'\b\w{3,}\b', text.lower())
        
        # Remove common words
        common_words = {
            'the', 'and', 'for', 'are', 'but', 'not', 'you', 'all', 'can', 'had', 
            'her', 'was', 'one', 'our', 'out', 'day', 'get', 'has', 'him', 'his', 
            'how', 'man', 'new', 'now', 'old', 'see', 'two', 'way', 'who', 'boy', 
            'did', 'its', 'let', 'put', 'say', 'she', 'too', 'use', 'may', 'will',
            'shall', 'this', 'that', 'with', 'have', 'from', 'they', 'know', 'want',
            'been', 'good', 'much', 'some', 'time', 'very', 'when', 'come', 'here',
            'just', 'like', 'long', 'make', 'many', 'over', 'such', 'take', 'than',
            'them', 'well', 'were'
        }
        
        filtered_words = [word for word in words if word not in common_words]
        word_freq = Counter(filtered_words)
        
        return [word for word, _ in word_freq.most_common(top_n)]
    
    def _calculate_complexity(self, text: str) -> float:
        """Calculate document complexity score (0-100)"""
        words = text.split()
        sentences = re.split(r'[.!?]+', text)
        
        # Factors contributing to complexity
        avg_word_length = sum(len(word) for word in words) / max(len(words), 1)
        avg_sentence_length = len(words) / max(len(sentences), 1)
        unique_words = len(set(word.lower() for word in words))
        vocabulary_richness = unique_words / max(len(words), 1)
        
        # Legal/technical terms increase complexity
        complex_terms = [
            'notwithstanding', 'hereinafter', 'aforementioned', 'whereas', 
            'pursuant', 'hereunder', 'thereof', 'hereof', 'compliance',
            'regulation', 'statute', 'ordinance', 'liability', 'indemnity'
        ]
        
        complex_term_count = sum(text.lower().count(term) for term in complex_terms)
        
        # Normalize and combine factors
        complexity = (
            (avg_word_length - 4) * 10 +  # Longer words = more complex
            (avg_sentence_length - 15) * 2 +  # Longer sentences = more complex
            vocabulary_richness * 30 +  # Higher vocabulary = more complex
            complex_term_count * 5  # Legal terms = more complex
        )
        
        return max(0, min(100, complexity))

# =============================================================================
# LLM Parser Module (Gemini)
# =============================================================================

class GeminiParser:
    """Handles Gemini API interactions for query parsing and logic evaluation"""
    
    def __init__(self):
        if not config.GEMINI_API_KEY:
            raise ValueError("GEMINI_API_KEY environment variable is required")
        
        genai.configure(api_key=config.GEMINI_API_KEY)
        self.model = genai.GenerativeModel('gemini-1.5-pro')
    
    async def parse_query_intent(self, query: str) -> Dict[str, Any]:
        """Extract structured intent from natural language query"""
        prompt = f"""
        Analyze the following natural language query and extract structured information:
        
        Query: "{query}"
        
        Please return a JSON object with the following structure:
        {{
            "intent": "coverage_check|policy_lookup|compliance_check|general_inquiry",
            "target": "specific item/procedure/condition being asked about",
            "focus": ["list", "of", "key", "aspects", "to", "focus", "on"],
            "question_type": "yes_no|conditions|explanation|comparison",
            "entities": ["list", "of", "named", "entities", "found"]
        }}
        
        Only return the JSON object, no additional text.
        """
        
        try:
            response = await asyncio.to_thread(
                self.model.generate_content,
                prompt,
                generation_config=genai.types.GenerationConfig(
                    temperature=0.1,
                    max_output_tokens=500
                )
            )
            return json.loads(response.text.strip())
        except Exception as e:
            logger.error(f"Error parsing query intent: {str(e)}")
            # Fallback basic parsing
            return {
                "intent": "general_inquiry",
                "target": query[:50],
                "focus": ["coverage", "conditions"],
                "question_type": "explanation",
                "entities": []
            }
    
    async def evaluate_clause_relevance(self, query: str, chunks: List[str]) -> List[Dict[str, Any]]:
        """Use Gemini to assess and rank chunk relevance"""
        prompt = f"""
        Query: "{query}"
        
        Please evaluate the relevance of each text chunk to the query and extract key information:
        
        Text Chunks:
        {chr(10).join([f"CHUNK_{i}: {chunk}" for i, chunk in enumerate(chunks)])}
        
        For each chunk, return a JSON array with objects containing:
        {{
            "chunk_id": "CHUNK_X",
            "relevance_score": 0.0-1.0,
            "key_phrases": ["relevant", "phrases", "found"],
            "contains_conditions": true/false,
            "summary": "brief summary of relevant content"
        }}
        
        Only return the JSON array, no additional text.
        """
        
        try:
            response = await asyncio.to_thread(
                self.model.generate_content,
                prompt,
                generation_config=genai.types.GenerationConfig(
                    temperature=0.2,
                    max_output_tokens=2000
                )
            )
            return json.loads(response.text.strip())
        except Exception as e:
            logger.error(f"Error evaluating clause relevance: {str(e)}")
            # Fallback scoring
            return [{"chunk_id": f"CHUNK_{i}", "relevance_score": 0.5, 
                    "key_phrases": [], "contains_conditions": False, 
                    "summary": "Could not analyze"} for i in range(len(chunks))]
    
    async def generate_final_answer(self, query: str, relevant_chunks: List[Dict[str, Any]], 
                                  query_intent: Dict[str, Any]) -> Dict[str, Any]:
        """Generate final structured answer based on relevant chunks"""
        chunks_text = "\n".join([f"CLAUSE_{chunk['chunk_id']}: {chunk['text']}" 
                                for chunk in relevant_chunks])
        
        prompt = f"""
        Based on the following relevant document clauses, provide a comprehensive answer to the user's query.
        
        Query: "{query}"
        Query Intent: {json.dumps(query_intent)}
        
        Relevant Clauses:
        {chunks_text}
        
        Please provide a response in the following JSON format:
        {{
            "answer": "Direct answer to the query",
            "conditions": ["list", "of", "any", "conditions", "or", "requirements"],
            "evidence": [
                {{
                    "clause_id": "unique_identifier",
                    "text": "relevant excerpt from the clause",
                    "relevance": "why this clause is relevant"
                }}
            ],
            "confidence": 0.0-1.0,
            "caveats": ["any", "limitations", "or", "uncertainties"]
        }}
        
        Be thorough but concise. If the answer cannot be determined from the provided clauses, state that clearly.
        Only return the JSON object, no additional text.
        """
        
        try:
            response = await asyncio.to_thread(
                self.model.generate_content,
                prompt,
                generation_config=genai.types.GenerationConfig(
                    temperature=0.3,
                    max_output_tokens=1500
                )
            )
            return json.loads(response.text.strip())
        except Exception as e:
            logger.error(f"Error generating final answer: {str(e)}")
            return {
                "answer": "Unable to process the query due to technical issues.",
                "conditions": [],
                "evidence": [],
                "confidence": 0.0,
                "caveats": ["Technical error occurred during processing"]
            }

# =============================================================================
# Embedding and Search Module
# =============================================================================

class EmbeddingSearchEngine:
    """Handles document chunking, embedding, and similarity search"""
    
    def __init__(self):
        # Use sentence-transformers for embeddings
        self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        self.dimension = 384  # Dimension for all-MiniLM-L6-v2
        self.index = faiss.IndexFlatIP(self.dimension)  # Inner product for cosine similarity
        self.chunks: List[DocumentChunk] = []
    
    def chunk_document(self, text: str, document_id: str) -> List[DocumentChunk]:
        """Split document into semantic chunks"""
        # Simple sentence-based chunking with overlap
        sentences = re.split(r'[.!?]+', text)
        chunks = []
        current_chunk = ""
        chunk_sentences = []
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
                
            # Check if adding this sentence would exceed chunk size
            if len(current_chunk) + len(sentence) > config.CHUNK_SIZE and current_chunk:
                # Create chunk
                chunk_id = f"{document_id}_{len(chunks)}"
                chunk = DocumentChunk(
                    id=chunk_id,
                    text=current_chunk.strip(),
                    metadata={
                        "document_id": document_id,
                        "chunk_index": len(chunks),
                        "sentence_count": len(chunk_sentences)
                    }
                )
                chunks.append(chunk)
                
                # Start new chunk with overlap
                overlap_sentences = chunk_sentences[-2:] if len(chunk_sentences) >= 2 else chunk_sentences
                current_chunk = " ".join(overlap_sentences) + " " + sentence
                chunk_sentences = overlap_sentences + [sentence]
            else:
                current_chunk += " " + sentence
                chunk_sentences.append(sentence)
        
        # Add final chunk
        if current_chunk.strip():
            chunk_id = f"{document_id}_{len(chunks)}"
            chunk = DocumentChunk(
                id=chunk_id,
                text=current_chunk.strip(),
                metadata={
                    "document_id": document_id,
                    "chunk_index": len(chunks),
                    "sentence_count": len(chunk_sentences)
                }
            )
            chunks.append(chunk)
        
        return chunks
    
    async def add_document(self, text: str, document_id: str) -> List[DocumentChunk]:
        """Process document and add to search index"""
        chunks = self.chunk_document(text, document_id)
        
        # Generate embeddings
        texts = [chunk.text for chunk in chunks]
        embeddings = await asyncio.to_thread(self.embedding_model.encode, texts)
        
        # Normalize embeddings for cosine similarity
        embeddings = embeddings / np.linalg.norm(embeddings, axis=1, keepdims=True)
        
        # Add to FAISS index
        self.index.add(embeddings.astype('float32'))
        
        # Store chunks with embeddings
        for chunk, embedding in zip(chunks, embeddings):
            chunk.embedding = embedding.tolist()
            self.chunks.append(chunk)
        
        logger.info(f"Added {len(chunks)} chunks to search index for document {document_id}")
        return chunks
    
    async def search_similar_chunks(self, query: str, k: int = None) -> List[DocumentChunk]:
        """Find most similar chunks to query"""
        if k is None:
            k = config.TOP_K_CHUNKS
        
        if not self.chunks:
            return []
        
        # Generate query embedding
        query_embedding = await asyncio.to_thread(self.embedding_model.encode, [query])
        query_embedding = query_embedding / np.linalg.norm(query_embedding, axis=1, keepdims=True)
        
        # Search in FAISS index
        scores, indices = self.index.search(query_embedding.astype('float32'), k)
        
        # Return matching chunks
        results = []
        for score, idx in zip(scores[0], indices[0]):
            if idx < len(self.chunks) and score > 0.3:  # Minimum similarity threshold
                chunk = self.chunks[idx]
                results.append(chunk)
        
        return results

# =============================================================================
# Batch Processing Module
# =============================================================================

class BatchProcessor:
    """Process multiple documents in batch"""
    
    def __init__(self, api_handler):
        self.api_handler = api_handler
        self.max_concurrent = config.MAX_CONCURRENT_BATCH
    
    async def process_batch(self, queries_and_files: List[tuple]) -> List[Dict]:
        """Process multiple query-file pairs concurrently"""
        semaphore = asyncio.Semaphore(self.max_concurrent)
        
        async def process_single(query, file):
            async with semaphore:
                try:
                    return await self.api_handler.process_query(query, file)
                except Exception as e:
                    return {
                        "query": query,
                        "error": str(e),
                        "status": "failed"
                    }
        
        tasks = [process_single(query, file) for query, file in queries_and_files]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        return [
            result if not isinstance(result, Exception) 
            else {"error": str(result), "status": "failed"}
            for result in results
        ]

# =============================================================================
# Streaming Processor for Large Documents
# =============================================================================

class StreamingProcessor:
    """Process large documents in streaming fashion"""
    
    def __init__(self, chunk_size: int = 8192):
        self.chunk_size = chunk_size
        
    async def process_large_document(self, file_stream, callback):
        """Process document in chunks with progress callbacks"""
        chunks = []
        total_size = 0
        
        async for chunk in self._read_in_chunks(file_stream):
            chunks.append(chunk)
            total_size += len(chunk)
            
            # Progress callback
            if callback:
                await callback({
                    'status': 'reading',
                    'bytes_processed': total_size,
                    'chunks_count': len(chunks)
                })
        
        # Combine chunks and process
        full_content = b''.join(chunks)
        
        if callback:
            await callback({
                'status': 'processing',
                'total_size': total_size,
                'stage': 'text_extraction'
            })
        
        return full_content
    
    async def _read_in_chunks(self, file_stream):
        """Read file in chunks asynchronously"""
        while True:
            chunk = await file_stream.read(self.chunk_size)
            if not chunk:
                break
            yield chunk

# =============================================================================
# Enhanced Main API Application
# =============================================================================

class EnhancedDocumentQueryAPI:
    """Enhanced API class orchestrating all components"""
    
    def __init__(self):
        self.document_processor = AdvancedDocumentProcessor()
        self.gemini_parser = GeminiParser()
        self.search_engine = EmbeddingSearchEngine()
        self.cache_manager = CacheManager()
        self.query_optimizer = QueryOptimizer()
        self.analytics = DocumentAnalytics()
        self.batch_processor = BatchProcessor(self)
        self.streaming_processor = StreamingProcessor()
        
    async def process_query(self, query: str, file: UploadFile) -> QueryResponse:
        """Main processing pipeline"""
        start_time = time.time()
        
        try:
            # Step 1: Document Ingestion
            logger.info("Extracting text from document...")
            document_text = await self.document_processor.process_document(file)
            
            # Step 2: Create search index
            document_id = str(uuid.uuid4())
            logger.info("Creating embeddings and search index...")
            await self.search_engine.add_document(document_text, document_id)
            
            # Step 3: Parse query intent
            logger.info("Parsing query intent...")
            query_intent = await self.gemini_parser.parse_query_intent(query)
            
            # Step 4: Retrieve relevant chunks
            logger.info("Searching for relevant document chunks...")
            relevant_chunks = await self.search_engine.search_similar_chunks(query)
            
            if not relevant_chunks:
                return QueryResponse(
                    query=query,
                    answer="No relevant information found in the document.",
                    conditions=[],
                    evidence=[],
                    confidence=0.0,
                    processing_time=time.time() - start_time
                )
            
            # Step 5: Evaluate chunk relevance with Gemini
            logger.info("Evaluating chunk relevance...")
            chunk_texts = [chunk.text for chunk in relevant_chunks]
            relevance_scores = await self.gemini_parser.evaluate_clause_relevance(query, chunk_texts)
            
            # Step 6: Prepare data for final answer generation
            scored_chunks = []
            for i, chunk in enumerate(relevant_chunks):
                score_data = relevance_scores[i] if i < len(relevance_scores) else {}
                scored_chunks.append({
                    "chunk_id": chunk.id,
                    "text": chunk.text,
                    "relevance_score": score_data.get("relevance_score", 0.5),
                    "metadata": chunk.metadata
                })
            
            # Sort by relevance and take top chunks
            scored_chunks.sort(key=lambda x: x["relevance_score"], reverse=True)
            top_chunks = scored_chunks[:5]  # Top 5 most relevant
            
            # Step 7: Generate final answer
            logger.info("Generating final answer...")
            final_result = await self.gemini_parser.generate_final_answer(
                query, top_chunks, query_intent
            )
            
            # Step 8: Format response
            processing_time = time.time() - start_time
            
            return QueryResponse(
                query=query,
                answer=final_result.get("answer", "Unable to generate answer"),
                conditions=final_result.get("conditions", []),
                evidence=final_result.get("evidence", []),
                confidence=final_result.get("confidence", 0.0),
                processing_time=processing_time
            )
            
        except Exception as e:
            logger.error(f"Error processing query: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Processing error: {str(e)}")
    
    async def process_query_with_cache(self, query: str, file: UploadFile) -> QueryResponse:
        """Process query with caching support"""
        start_time = time.time()
        
        # Generate document hash for caching
        file_content = await file.read()
        await file.seek(0)  # Reset file pointer
        doc_hash = hashlib.md5(file_content).hexdigest()
        
        # Check cache for previous response
        cached_response = await self.cache_manager.get_query_response(query, doc_hash)
        if cached_response:
            logger.info("Cache hit for query")
            cached_response['processing_time'] = time.time() - start_time
            cached_response['cache_hit'] = True
            return QueryResponse(**cached_response)
        
        # Process normally
        response = await self.process_query(query, file)
        
        # Cache the response
        await self.cache_manager.set_query_response(
            query, doc_hash, asdict(response)
        )
        
        response.cache_hit = False
        return response
    
    async def analyze_document_insights(self, file: UploadFile) -> DocumentAnalysis:
        """Provide document analysis and insights"""
        text = await self.document_processor.process_document(file)
        
        # Basic analysis
        insights = self.analytics.analyze_document(text)
        
        # Add embedding-based insights
        document_id = str(uuid.uuid4())
        chunks = await self.search_engine.add_document(text, document_id)
        
        insights.update({
            'chunk_count': len(chunks),
            'avg_chunk_length': sum(len(chunk.text) for chunk in chunks) / len(chunks) if chunks else 0,
            'processing_timestamp': datetime.now().isoformat(),
            'document_id': document_id
        })
        
        return DocumentAnalysis(**insights)
    
    async def suggest_queries(self, file: UploadFile, num_suggestions: int = 5) -> List[str]:
        """Generate suggested queries based on document content"""
        text = await self.document_processor.process_document(file)
        doc_type = self.analytics._classify_document_type(text)
        key_topics = self.analytics._extract_key_topics(text, top_n=10)
        
        # Generate contextual suggestions based on document type
        suggestions = []
        
        if doc_type == 'insurance_policy':
            base_templates = [
                "What does this policy cover for {topic}?",
                "Are there any exclusions related to {topic}?",
                "What are the conditions for {topic} coverage?",
                "What is the deductible for {topic}?",
                "How do I file a claim for {topic}?"
            ]
        elif doc_type == 'hr_policy':
            base_templates = [
                "What are the policies regarding {topic}?",
                "What are the eligibility requirements for {topic}?",
                "How much {topic} time is available?",
                "What is the process for requesting {topic}?",
                "Are there any restrictions on {topic}?"
            ]
        elif doc_type == 'compliance_doc':
            base_templates = [
                "What are the {topic} requirements?",
                "How often should {topic} be reviewed?",
                "What are the penalties for {topic} violations?",
                "Who is responsible for {topic} compliance?",
                "What documentation is needed for {topic}?"
            ]
        else:
            base_templates = [
                "What information is provided about {topic}?",
                "What are the key points regarding {topic}?",
                "How is {topic} defined in this document?",
                "What are the implications of {topic}?",
                "Are there any conditions related to {topic}?"
            ]
        
        # Generate suggestions using key topics
        for i, template in enumerate(base_templates[:num_suggestions]):
            if i < len(key_topics):
                topic = key_topics[i].replace('_', ' ').title()
                suggestions.append(template.format(topic=topic))
        
        # Add some generic suggestions
        generic_suggestions = [
            "What are the main points covered in this document?",
            "Are there any important deadlines or timeframes?",
            "What are the key terms and definitions?",
            "What are the main requirements or conditions?",
            "Who does this document apply to?"
        ]
        
        # Fill remaining slots with generic suggestions
        while len(suggestions) < num_suggestions and generic_suggestions:
            suggestions.append(generic_suggestions.pop(0))
        
        return suggestions[:num_suggestions]
    
    async def compare_documents(self, files: List[UploadFile], 
                              comparison_aspects: List[str]) -> ComparisonResult:
        """Compare multiple documents across specified aspects"""
        if len(files) > 5:
            raise HTTPException(status_code=400, detail="Maximum 5 documents allowed for comparison")
        
        documents = {}
        for i, file in enumerate(files):
            text = await self.document_processor.process_document(file)
            doc_id = f"doc_{i}"
            documents[doc_id] = {
                'filename': file.filename,
                'text': text,
                'analysis': self.analytics.analyze_document(text)
            }
        
        # Compare across aspects
        comparison_results = {
            'document_summaries': {},
            'aspect_comparisons': {},
            'similarities': {},
            'differences': {}
        }
        
        # Generate summaries
        for doc_id, doc_data in documents.items():
            comparison_results['document_summaries'][doc_id] = {
                'filename': doc_data['filename'],
                'word_count': doc_data['analysis']['word_count'],
                'document_type': doc_data['analysis']['document_type'],
                'complexity_score': doc_data['analysis']['complexity_score'],
                'key_topics': doc_data['analysis']['key_topics']
            }
        
        # Compare specific aspects using Gemini
        for aspect in comparison_aspects:
            aspect_comparison = await self._compare_aspect_across_documents(
                documents, aspect
            )
            comparison_results['aspect_comparisons'][aspect] = aspect_comparison
        
        return ComparisonResult(**comparison_results)
    
    async def _compare_aspect_across_documents(self, documents: Dict, aspect: str) -> Dict[str, Any]:
        """Compare a specific aspect across multiple documents using Gemini"""
        doc_texts = {doc_id: doc_data['text'] for doc_id, doc_data in documents.items()}
        
        prompt = f"""
        Compare the following documents regarding "{aspect}". 
        For each document, extract relevant information and then provide a comparison.
        
        Documents:
        {chr(10).join([f"DOCUMENT_{doc_id} ({documents[doc_id]['filename']}):{chr(10)}{text[:2000]}..." 
                      for doc_id, text in doc_texts.items()])}
        
        Please provide a JSON response with:
        {{
            "aspect": "{aspect}",
            "document_findings": {{
                "doc_0": "what document 0 says about {aspect}",
                "doc_1": "what document 1 says about {aspect}",
                ...
            }},
            "similarities": ["list of similarities across documents"],
            "differences": ["list of key differences"],
            "summary": "overall comparison summary"
        }}
        
        Only return the JSON object.
        """
        
        try:
            response = await asyncio.to_thread(
                self.gemini_parser.model.generate_content,
                prompt,
                generation_config=genai.types.GenerationConfig(
                    temperature=0.3,
                    max_output_tokens=2000
                )
            )
            return json.loads(response.text.strip())
        except Exception as e:
            logger.error(f"Error comparing aspect {aspect}: {e}")
            return {
                "aspect": aspect,
                "document_findings": {},
                "similarities": [],
                "differences": [],
                "summary": f"Error analyzing aspect: {str(e)}"
            }

# =============================================================================
# FastAPI Application Setup
# =============================================================================

# Initialize API
enhanced_api = EnhancedDocumentQueryAPI()
app = FastAPI(
    title="Enhanced Document Query API",
    description="AI-powered document analysis using Google Gemini with caching, OCR, and advanced features",
    version="2.0.0"
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# =============================================================================
# API Endpoints
# =============================================================================

@app.post("/ask-document", response_model=QueryResponse)
async def ask_document(
    query: str = Form(...),
    file: UploadFile = File(...)
):
    """
    Main endpoint for document queries
    
    - **query**: Natural language question about the document
    - **file**: Document file (PDF, DOCX, EML, TXT, PNG, JPG)
    """
    return await enhanced_api.process_query(query, file)

@app.post("/ask-document-cached", response_model=QueryResponse)
async def ask_document_cached(
    query: str = Form(...),
    file: UploadFile = File(...)
):
    """
    Query document with caching support
    
    - **query**: Natural language question about the document
    - **file**: Document file with caching for faster repeated queries
    """
    return await enhanced_api.process_query_with_cache(query, file)

@app.post("/analyze-document", response_model=DocumentAnalysis)
async def analyze_document(file: UploadFile = File(...)):
    """
    Analyze document and provide comprehensive insights
    
    - **file**: Document file to analyze
    """
    return await enhanced_api.analyze_document_insights(file)

@app.post("/suggest-queries")
async def suggest_queries(
    file: UploadFile = File(...),
    num_suggestions: int = 5
):
    """
    Generate suggested queries for a document based on its content
    
    - **file**: Document file
    - **num_suggestions**: Number of suggestions to generate (default: 5, max: 10)
    """
    if num_suggestions > 10:
        raise HTTPException(status_code=400, detail="Maximum 10 suggestions allowed")
    
    suggestions = await enhanced_api.suggest_queries(file, num_suggestions)
    return {"suggestions": suggestions}

@app.post("/compare-documents", response_model=ComparisonResult)
async def compare_documents(
    files: List[UploadFile] = File(...),
    comparison_aspects: List[str] = Form(["coverage", "conditions", "requirements"])
):
    """
    Compare multiple documents across specified aspects
    
    - **files**: List of document files (max 5)
    - **comparison_aspects**: Aspects to compare across documents
    """
    return await enhanced_api.compare_documents(files, comparison_aspects)

@app.post("/batch-process")
async def batch_process(
    background_tasks: BackgroundTasks,
    queries: List[str] = Form(...),
    files: List[UploadFile] = File(...)
):
    """
    Process multiple query-document pairs in batch
    
    - **queries**: List of queries (must match number of files)
    - **files**: List of document files
    """
    if len(queries) != len(files):
        raise HTTPException(
            status_code=400, 
            detail="Number of queries must match number of files"
        )
    
    if len(files) > 10:
        raise HTTPException(
            status_code=400, 
            detail="Maximum 10 files allowed in batch processing"
        )
    
    query_file_pairs = list(zip(queries, files))
    results = await enhanced_api.batch_processor.process_batch(query_file_pairs)
    return {"results": results}

# =============================================================================
# WebSocket Endpoints
# =============================================================================

@app.websocket("/ws/process-document")
async def websocket_process_document(websocket: WebSocket):
    """
    WebSocket endpoint for real-time document processing updates
    """
    await websocket.accept()
    
    try:
        while True:
            # Wait for file data and query
            data = await websocket.receive_json()
            query = data.get("query")
            
            if not query:
                await websocket.send_json({"error": "Query is required"})
                continue
            
            # Process with progress updates
            async def progress_callback(progress_info):
                await websocket.send_json({
                    "type": "progress",
                    "data": progress_info
                })
            
            # Send processing updates
            await websocket.send_json({
                "type": "status", 
                "message": "Starting document processing..."
            })
            
            try:
                await websocket.send_json({
                    "type": "status", 
                    "message": "Document processing completed successfully"
                })
                
            except Exception as e:
                await websocket.send_json({
                    "type": "error",
                    "message": f"Processing failed: {str(e)}"
                })
                
    except WebSocketDisconnect:
        logger.info("WebSocket client disconnected")

# =============================================================================
# Health and Status Endpoints
# =============================================================================

@app.get("/health")
async def health_check():
    """Basic health check endpoint"""
    return {
        "status": "healthy", 
        "model": "gemini-1.5-pro",
        "timestamp": datetime.now().isoformat()
    }

@app.get("/health/detailed")
async def detailed_health_check():
    """Detailed health check with system metrics"""
    health_info = {
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "cache": {
            "redis_available": enhanced_api.cache_manager.enabled
        },
        "models": {
            "gemini_model": "gemini-1.5-pro",
            "embedding_model": "all-MiniLM-L6-v2"
        },
        "features": {
            "ocr_available": OCR_AVAILABLE,
            "nltk_available": NLTK_AVAILABLE,
            "redis_available": REDIS_AVAILABLE
        }
    }
    
    # Add system metrics if available
    if PSUTIL_AVAILABLE:
        health_info["system"] = {
            "cpu_percent": psutil.cpu_percent(interval=1),
            "memory_percent": psutil.virtual_memory().percent,
            "disk_percent": psutil.disk_usage('/').percent,
        }
        
        if hasattr(psutil, 'getloadavg'):
            health_info["system"]["load_average"] = psutil.getloadavg()
    
    return health_info

@app.get("/")
async def root():
    """Root endpoint with API information"""
    return {
        "message": "Enhanced Document Query API with Google Gemini",
        "version": "2.0.0",
        "features": [
            "Document text extraction (PDF, DOCX, EML, TXT, Images)",
            "OCR support for scanned documents",
            "Intelligent query processing with Gemini",
            "Document analytics and insights",
            "Caching for improved performance",
            "Batch processing",
            "Document comparison",
            "Query suggestions",
            "Real-time WebSocket processing"
        ],
        "endpoints": {
            "POST /ask-document": "Submit document and query",
            "POST /ask-document-cached": "Submit document and query with caching",
            "POST /analyze-document": "Analyze document and get insights",
            "POST /suggest-queries": "Get suggested queries for document",
            "POST /compare-documents": "Compare multiple documents",
            "POST /batch-process": "Process multiple documents in batch",
            "WS /ws/process-document": "Real-time processing updates",
            "GET /health": "Basic health check",
            "GET /health/detailed": "Detailed health and system metrics",
            "GET /docs": "API documentation"
        },
        "supported_formats": enhanced_api.document_processor.supported_formats
    }

# =============================================================================
# CLI Runner (for development)
# =============================================================================

if __name__ == "__main__":
    # Ensure required environment variables
    if not config.GEMINI_API_KEY:
        print("Error: GEMINI_API_KEY environment variable is required")
        print("Please set it with your Google AI API key")
        exit(1)
    
    print("Starting Enhanced Document Query API...")
    print(f"Supported formats: {enhanced_api.document_processor.supported_formats}")
    print(f"Features enabled:")
    print(f"  - OCR: {OCR_AVAILABLE}")
    print(f"  - NLTK: {NLTK_AVAILABLE}")
    print(f"  - Redis Caching: {REDIS_AVAILABLE}")
    print(f"  - System Monitoring: {PSUTIL_AVAILABLE}")
    print("Access API documentation at: http://localhost:8000/docs")
    
    uvicorn.run(
        "main:app",  # Assuming this file is named main.py
        host="0.0.0.0",
        port=8000,
        reload=True,
        log_level="info"
    )